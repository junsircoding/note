"""python-docx 基本使用"""

from docx.document import Document
from docx import Document as document_func
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.enum.text import (
    WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT,
    WD_TAB_LEADER, WD_UNDERLINE
)
from docx.enum.section import (
    WD_SECTION_START, WD_ORIENTATION
)

def document_with_paragraph(document: Document):
    # 添加标题
    document.add_heading('0级标题', 0) # 标题级别, 范围 0-9

    # 加粗和斜体样式
    p = document.add_paragraph('测试文本')
    p.add_run('加粗文本').bold = True
    p.add_run('普通文本')
    p.add_run('斜体文本。').italic = True

    document.add_heading('1级标题', level=1)
    # 引用样式
    document.add_paragraph('明显引用', style='Intense Quote')
    
    # 列表
    document.add_paragraph(
        '子弹列表第一项', style='List Bullet'
    )
    document.add_paragraph(
        '子弹列表第二项', style='List Bullet'
    )
    document.add_paragraph(
        '数字列表第一项', style='List Number'
    )
    document.add_paragraph(
        '数字列表第二项', style='List Number'
    )
    
    # 添加图片
    document.add_picture('monty-truth.jpeg', width=Inches(1.25))

    # 添加表格
    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam')
    )

    table = document.add_table(rows=1, cols=3)
    table.style = "Light Shading Accent 1"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc

    # 添加分页符
    document.add_page_break()

    # 文本左缩进
    paragraph = document.add_paragraph("文本左缩进")
    paragraph.add_run(" 0.5 Inch", "Emphasis")
    paragraph_format = paragraph.paragraph_format
    paragraph_format.left_indent = Inches(0.5)

    # 文本右缩进
    paragraph = document.add_paragraph("文本右缩进为")
    paragraph.add_run(" 24 Pt", "Emphasis")
    paragraph_format = paragraph.paragraph_format
    paragraph_format.left_indent = Inches(0.5)

    # 文本居中
    paragraph = document.add_paragraph("文本居中")
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 字体
    paragraph = document.add_paragraph()
    run = paragraph.add_run("Calibri 字体, 字号为 12 Pt")
    font = run.font
    font.name = "Calibri"
    font.size = Pt(12)

    # 字体高亮, 1-7
    paragraph = document.add_paragraph()
    for color_idx, color_name in {
        1: "黑色", 2: "蓝色", 3: "青色",
        4: "绿色", 5: "粉色", 6: "红色",
        7: "黄色",
    }.items():
        run = paragraph.add_run(f"字体高亮-[{color_idx}]-[{color_name}]"+"\n")
        font = run.font
        if color_idx == 1:
            font.color.rgb = RGBColor(255, 255, 255)
        font.highlight_color = color_idx

    # 下划线
    paragraph = document.add_paragraph()
    run = paragraph.add_run("下划线")
    font = run.font
    font.underline = True
    font.underline = WD_UNDERLINE.SINGLE

    # 字体颜色
    paragraph = document.add_paragraph()
    run = paragraph.add_run("字体颜色")
    font = run.font
    font.color.rgb = RGBColor(0x42, 0x24, 0xE9)

    # 字符底纹
    paragraph = document.add_paragraph()
    run = paragraph.add_run("字符底纹")
    font = run.font
    font.color.theme_color = MSO_THEME_COLOR_INDEX.ACCENT_1

    # 文本悬挂缩进
    paragraph = document.add_paragraph("文本悬挂缩进 -0.25")
    paragraph_format = paragraph.paragraph_format
    paragraph_format.first_line_indent = Inches(-0.25)

    # 制表符
    paragraph = document.add_paragraph("制表符")
    paragraph_format = paragraph.paragraph_format
    tab_stops = paragraph_format.tab_stops
    tab_stop = tab_stops.add_tab_stop(Inches(1.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)

    # 段落格式
    paragraph = document.add_paragraph("段落格式，段前 18 Pt, 段后 12 Pt")
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(18)
    paragraph_format.space_after = Pt(12)

    # 行空白
    paragraph = document.add_paragraph("行空白")
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing = Pt(18)
    paragraph_format.line_spacing = 1.75

    # 页码
    paragraph = document.add_paragraph("页码")
    paragraph_format = paragraph.paragraph_format
    paragraph_format.keep_with_next = True
    paragraph_format.page_break_before = False

    # 保存文档
    document.save('demo.docx')

def document_with_section(document: Document):
    current_section = document.sections[-1]  # last section in document
    print(current_section.start_type)
    new_section = document.add_section(WD_SECTION_START.ODD_PAGE)
    print(new_section.start_type)

    section = current_section
    print("%s, %s, %s" % (section.orientation, section.page_width, section.page_height))
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height
    print("%s, %s, %s" % (section.orientation, section.page_width, section.page_height))

    # 保存文档
    document.save("demo2.docx")

if __name__ == "__main__":
    document = document_func()
    document_with_paragraph(document)
